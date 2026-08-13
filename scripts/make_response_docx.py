#!/usr/bin/env python3
"""Render the response-to-reviewers markdown into the .docx IEEE Access asks for.

The decision letter wants a document with, for each comment, the reviewer's concern, our
response, and the action taken. We keep the authored content in markdown because it is
diffable and reviewable; this script is only the last mile to the format the portal
accepts.

    python scripts/make_response_docx.py

Writes Review/Response_to_Reviewers_Access-2026-28984.docx
"""
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
SRC = ROOT / "Review" / "RESPONSE_TO_REVIEWERS.md"
DEST = ROOT / "Review" / "Response_to_Reviewers_Access-2026-28984.docx"

ACCENT = RGBColor(0x1F, 0x4E, 0x79)


def add_runs(par, text):
    """Render the small markdown subset the letter uses: bold, italic, and inline code."""
    for chunk in re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`|\$[^$]+\$)', text):
        if not chunk:
            continue
        if chunk.startswith("**") and chunk.endswith("**"):
            par.add_run(chunk[2:-2]).bold = True
        elif chunk.startswith("*") and chunk.endswith("*") and len(chunk) > 2:
            par.add_run(chunk[1:-1]).italic = True
        elif chunk.startswith("`") and chunk.endswith("`"):
            r = par.add_run(chunk[1:-1])
            r.font.name = "Consolas"
            r.font.size = Pt(9.5)
        elif chunk.startswith("$") and chunk.endswith("$"):
            # inline maths in a Word document reads better as plain text than as TeX
            r = par.add_run(chunk[1:-1].replace("\\", "").replace("{", "").replace("}", ""))
            r.italic = True
        else:
            par.add_run(chunk)


def main():
    md = SRC.read_text(encoding="utf-8")
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)
    style.paragraph_format.space_after = Pt(6)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("Response to Reviewers")
    r.bold = True
    r.font.size = Pt(16)
    r.font.color.rgb = ACCENT

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("Manuscript ID Access-2026-28984\n"
                    "AgentFairBench: Do LLM Agents Discriminate When They Act?")
    r.font.size = Pt(11)

    in_table, table_rows = False, []

    def flush_table():
        nonlocal in_table, table_rows
        if not table_rows:
            in_table = False
            return
        rows = [r for r in table_rows if not re.match(r'^\s*\|[\s:|-]+\|\s*$', r)]
        cells = [[c.strip() for c in r.strip().strip('|').split('|')] for r in rows]
        ncol = max(len(c) for c in cells)
        t = doc.add_table(rows=0, cols=ncol)
        t.style = "Light Grid Accent 1"
        for i, row in enumerate(cells):
            row = row + [""] * (ncol - len(row))
            wr = t.add_row().cells
            for j, cell in enumerate(row):
                wr[j].text = ""
                add_runs(wr[j].paragraphs[0], cell)
                if i == 0:
                    for p in wr[j].paragraphs:
                        for run in p.runs:
                            run.bold = True
        doc.add_paragraph()
        table_rows = []
        in_table = False

    for line in md.splitlines():
        s = line.rstrip()
        if s.startswith("|") and s.count("|") >= 2:
            in_table = True
            table_rows.append(s)
            continue
        if in_table:
            flush_table()

        if s.startswith("### "):
            p = doc.add_paragraph()
            r = p.add_run(s[4:])
            r.bold = True
            r.font.size = Pt(12)
            r.font.color.rgb = ACCENT
        elif s.startswith("## "):
            doc.add_page_break()
            p = doc.add_paragraph()
            r = p.add_run(s[3:])
            r.bold = True
            r.font.size = Pt(14)
            r.font.color.rgb = ACCENT
        elif s.startswith("# "):
            continue
        elif s.strip() in ("---", "***"):
            continue
        elif re.match(r'^\s*[-*]\s+', s):
            p = doc.add_paragraph(style="List Bullet")
            add_runs(p, re.sub(r'^\s*[-*]\s+', '', s))
        elif re.match(r'^\s*\d+\.\s+', s):
            p = doc.add_paragraph(style="List Number")
            add_runs(p, re.sub(r'^\s*\d+\.\s+', '', s))
        elif not s.strip():
            continue
        else:
            add_runs(doc.add_paragraph(), s)

    if in_table:
        flush_table()

    doc.save(DEST)
    words = len(md.split())
    print(f"wrote {DEST}")
    print(f"  {words} words, {len(doc.paragraphs)} paragraphs, {len(doc.tables)} tables")


if __name__ == "__main__":
    main()
